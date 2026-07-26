"""Classification gate: only Public documents may be uploaded; consent required."""

from __future__ import annotations

import io

import bcrypt
import pytest
from docx import Document
from fastapi.testclient import TestClient

from app import config
from app.auth import verify_credentials
from app.main import app
from app.services.document_classification import (
    CONSENT_POLICY,
    find_classification_marking,
)
from app.services.document_store import reset_document_store_for_tests
from app.services.token_store import reset_refresh_token_store_for_tests

client = TestClient(app)

TEST_SECRET = "test-secret-please-ignore-0123456789"
ALICE, ALICE_PW = "cls-alice", "alice-pass-xyz"


@pytest.fixture(autouse=True)
def _docs_env(monkeypatch, tmp_path):
    monkeypatch.setattr(config, "JWT_SECRET_KEY", TEST_SECRET)
    monkeypatch.setattr(config, "JWT_ALGORITHM", "HS256")
    monkeypatch.setattr(config, "JWT_ISSUER", "misa-intelligence-api")
    monkeypatch.setattr(config, "JWT_AUDIENCE", "misa-intelligence-api")
    monkeypatch.setattr(config, "ALLOW_PLAINTEXT_BOOTSTRAP", False)
    monkeypatch.setattr(config, "AUTH_DISABLED", False)
    monkeypatch.setattr(config, "API_USERNAME", "")
    monkeypatch.setattr(config, "API_PASSWORD", "")
    monkeypatch.setattr(config, "DOCUMENTS_ENABLED", True)
    monkeypatch.setattr(config, "DOCUMENTS_BACKEND", "memory")
    monkeypatch.setattr(config, "DOCUMENTS_ROOT", str(tmp_path / "docs"))
    monkeypatch.setattr(config, "DOCUMENTS_INGEST_DIR", str(tmp_path / "inbox"))
    monkeypatch.setattr(config, "MALWARE_SCAN_ENABLED", False)
    monkeypatch.setattr(config, "MALWARE_SCAN_BACKEND", "none")
    monkeypatch.setattr(config, "AUTH_USERS", {
        ALICE: bcrypt.hashpw(ALICE_PW.encode(), bcrypt.gensalt()).decode(),
    })
    monkeypatch.setattr(config, "AUTH_USER_ROLES", {ALICE: config.ROLE_ANALYST})
    reset_refresh_token_store_for_tests()
    reset_document_store_for_tests()
    app.dependency_overrides.pop(verify_credentials, None)
    yield


def _token() -> str:
    r = client.post("/api/v1/auth/login", json={"username": ALICE, "password": ALICE_PW})
    assert r.status_code == 200, r.text
    return r.json()["access_token"]


def _bearer() -> dict:
    return {"Authorization": f"Bearer {_token()}"}


def _upload(data: dict, filename="brief.txt", text="Acme opened an RHQ in Riyadh."):
    return client.post(
        "/api/v1/documents/upload",
        headers=_bearer(),
        files={"file": (filename, text.encode(), "text/plain")},
        data=data,
    )


def test_public_with_consent_is_accepted():
    r = _upload({"visibility": "private", "classification": "public", "consent": "true"})
    assert r.status_code == 200, r.text
    assert r.json()["status"] == "ready"


@pytest.mark.parametrize("label", ["restricted", "secret", "top_secret", "Top Secret"])
def test_non_public_labels_are_rejected(label):
    r = _upload({"visibility": "private", "classification": label, "consent": "true"})
    assert r.status_code == 403, r.text
    assert r.json()["error"]["code"] == "CLASSIFIED_DOCUMENT"


def test_missing_consent_is_rejected():
    r = _upload({"visibility": "private", "classification": "public"})
    assert r.status_code == 428, r.text
    assert r.json()["error"]["code"] == "CONSENT_REQUIRED"


def test_unknown_label_is_rejected():
    r = _upload({"visibility": "private", "classification": "internal", "consent": "true"})
    assert r.status_code == 400, r.text
    assert r.json()["error"]["code"] == "BAD_CLASSIFICATION"


@pytest.mark.parametrize("text", [
    "TOP SECRET\nQuarterly plans for the region.",
    "Top Secret briefing on energy policy.",
    "Classification: Restricted\nInternal distribution only.",
    "SECRET\nSource assessments follow.",
    "RESTRICTED handling caveats apply.",
    "CONFIDENTIAL\nBoard eyes only.",
    "سري للغاية\nتقرير داخلي",
])
def test_marked_content_is_rejected_and_not_stored(text):
    r = _upload(
        {"visibility": "private", "classification": "public", "consent": "true"},
        text=text,
    )
    assert r.status_code == 403, r.text
    assert r.json()["error"]["code"] == "CLASSIFIED_CONTENT"
    listed = client.get("/api/v1/documents", headers=_bearer())
    assert listed.json()["documents"] == []


def test_marked_filename_is_rejected():
    r = _upload(
        {"visibility": "private", "classification": "public", "consent": "true"},
        filename="secret-plan.txt",
    )
    assert r.status_code == 403, r.text
    assert r.json()["error"]["code"] == "CLASSIFIED_CONTENT"


def test_plain_prose_secret_is_not_a_marking():
    r = _upload(
        {"visibility": "private", "classification": "public", "consent": "true"},
        text="The recipe was a closely guarded trade secret of the founder.",
    )
    assert r.status_code == 200, r.text


def test_docx_header_marking_is_rejected():
    doc = Document()
    doc.add_paragraph("Routine market summary.")
    doc.sections[0].header.paragraphs[0].text = "TOP SECRET"
    buf = io.BytesIO()
    doc.save(buf)
    r = client.post(
        "/api/v1/documents/upload",
        headers=_bearer(),
        files={"file": ("summary.docx", buf.getvalue(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={"visibility": "private", "classification": "public", "consent": "true"},
    )
    assert r.status_code == 403, r.text
    assert r.json()["error"]["code"] == "CLASSIFIED_CONTENT"


def test_consent_policy_endpoint():
    r = client.get("/api/v1/documents/consent-policy", headers=_bearer())
    assert r.status_code == 200
    body = r.json()
    assert body["version"] == CONSENT_POLICY["version"]
    ids = [t["id"] for t in body["terms"]]
    assert {"classification", "content_standards", "personal_data",
            "rights", "processing"} <= set(ids)


def _labelled_docx(label: str) -> bytes:
    """Minimal DOCX carrying a Boldon James style security-label property."""
    doc = Document()
    doc.add_paragraph("Test file")
    buf = io.BytesIO()
    doc.save(buf)
    import zipfile
    out = io.BytesIO()
    with zipfile.ZipFile(buf) as src, zipfile.ZipFile(out, "w") as dst:
        for item in src.namelist():
            dst.writestr(item, src.read(item))
        dst.writestr(
            "docProps/custom.xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties"'
            ' xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">'
            '<property fmtid="{D5CDD505-2E9C-101B-9397-08002B2CF9AE}" pid="2"'
            f' name="bjDocumentSecurityLabel"><vt:lpwstr>{label}</vt:lpwstr></property>'
            "</Properties>",
        )
    return out.getvalue()


@pytest.mark.parametrize("label,accepted", [
    ("Public /  متاح", True),
    ("Restricted / مقيد", False),
    ("Secret /  سري", False),
    ("Top Secret / سري للغاية", False),
])
def test_classifier_tool_label_property(label, accepted):
    r = client.post(
        "/api/v1/documents/upload",
        headers=_bearer(),
        files={"file": ("labelled.docx", _labelled_docx(label),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={"visibility": "private", "consent": "true"},
    )
    if accepted:
        assert r.status_code == 200, r.text
    else:
        assert r.status_code == 403, r.text
        assert r.json()["error"]["code"] == "CLASSIFIED_CONTENT"


def test_marking_detector_unit():
    assert find_classification_marking("TOP SECRET plans") is not None
    assert find_classification_marking("top secret plans") is not None
    assert find_classification_marking("Classification: Secret") is not None
    assert find_classification_marking("a trade secret recipe") is None
    assert find_classification_marking("", "TOP-SECRET.pdf") is not None
    assert find_classification_marking("normal text", "briefing.pdf") is None
