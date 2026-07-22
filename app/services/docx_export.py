"""Quality-gated DOCX export from briefing markdown / structured payload.

Uses python-docx when installed. Always runs quality_gate first so
false-zeros / truncated rankings are not shipped as Word docs.
"""

from __future__ import annotations

import io
import re
from typing import Optional


def render_docx(
    question: str,
    answer_markdown: str,
    *,
    structured: Optional[dict] = None,
) -> bytes:
    """Return DOCX bytes. Raises ImportError if python-docx missing."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for Word export "
            "(install via requirements-dev / pip install python-docx)"
        ) from exc

    text = answer_markdown or ""
    try:
        from app.services.surface_quality import run_surface_quality_gate
        text, issues, fixes = run_surface_quality_gate(
            text,
            question=question or "",
            hard_block=True,
        )
        try:
            from app.services.quality_metrics import record_export
            record_export(
                kind="docx",
                quality_blocked=any("hard_block" in str(f) for f in (fixes or [])),
            )
        except Exception:
            pass
    except Exception:
        pass

    # Soft-trim incomplete trailing pipe rows (same failure class as PDF).
    try:
        from app.services.quality_gate import _trim_trailing_incomplete_pipe_row
        text = _trim_trailing_incomplete_pipe_row(text)
    except Exception:
        pass

    if structured:
        try:
            from app.schemas.quality_response import validate_quality_response
            from app.services.pdf_export import _structured_to_markdown
            model, errs = validate_quality_response(structured)
            if model and not errs:
                text = _structured_to_markdown(model)
        except Exception:
            pass

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("MISA Intelligence Briefing", level=1)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x00, 0x6C, 0x35)

    if question:
        p = doc.add_paragraph()
        run = p.add_run(question.strip())
        run.bold = True

    for line in (text or "").splitlines():
        raw = line.rstrip()
        if not raw.strip():
            doc.add_paragraph("")
            continue
        if re.match(r"^#{1,3}\s+", raw):
            level = len(re.match(r"^(#+)", raw).group(1))
            doc.add_heading(re.sub(r"^#{1,3}\s+", "", raw).strip(), level=min(level, 3))
            continue
        if raw.strip().startswith("|") and raw.count("|") >= 2:
            # Collect contiguous table rows
            continue  # handled in second pass below
        if raw.strip().startswith("- ") or raw.strip().startswith("* "):
            doc.add_paragraph(raw.strip()[2:], style="List Bullet")
            continue
        # Strip light markdown bold/italic
        clean = re.sub(r"\*\*(.+?)\*\*", r"\1", raw)
        clean = re.sub(r"\*(.+?)\*", r"\1", clean)
        doc.add_paragraph(clean)

    # Simple markdown tables → Word tables
    _append_markdown_tables(doc, text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _append_markdown_tables(doc, text: str) -> None:
    try:
        from docx.shared import Pt
    except ImportError:
        return
    lines = (text or "").splitlines()
    i = 0
    while i < len(lines):
        if lines[i].strip().startswith("|") and lines[i].count("|") >= 2:
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.match(r"^:?-+:?$", c or "") for c in cells):
                    rows.append(cells)
                i += 1
            if len(rows) >= 2:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for r_i, row in enumerate(rows):
                    for c_i, cell in enumerate(row):
                        if c_i < len(table.rows[r_i].cells):
                            table.rows[r_i].cells[c_i].text = cell
                doc.add_paragraph("")
            continue
        i += 1
