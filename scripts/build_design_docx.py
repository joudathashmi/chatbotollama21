"""
Build the interactive Word design document from DESIGN.md.

Output: docs/MISA_Chatbot_Design_Document.docx
        docs/diagrams/*.png

The .docx includes:
  - Cover page with title + version + date
  - Hyperlinked table of contents (Word's native TOC field — refresh in
    Word with F9 or right-click → Update Field)
  - All 16 sections + 2 appendices from DESIGN.md
  - 6 embedded architecture diagrams (auto-generated PNGs)
  - Word-native tables (not markdown) for tabular data
  - Heading styles so navigation pane works
  - Page numbers in footer
  - Hyperlinks back to source code paths and external compliance refs
"""

from __future__ import annotations

import os
import sys
from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


HERE = Path(__file__).parent
ROOT = HERE.parent
DOCS = ROOT / "docs"
DIAGRAMS = DOCS / "diagrams"
DIAGRAMS.mkdir(parents=True, exist_ok=True)


# ─── Brand palette ───────────────────────────────────────────────────
COLOR_PRIMARY = "#1a365d"      # MISA-ish navy
COLOR_ACCENT = "#0a73f0"       # interactive blue
COLOR_GREEN = "#2d6a4f"
COLOR_RED = "#c53030"
COLOR_AMBER = "#d97706"
COLOR_MUTED = "#6b7280"
COLOR_BG = "#f9fafb"
COLOR_BOX = "#ffffff"
COLOR_BOX_EDGE = "#d1d5db"


# ─── Diagram helpers ─────────────────────────────────────────────────

def _new_canvas(width: float = 12, height: float = 7, dpi: int = 150):
    fig, ax = plt.subplots(figsize=(width, height), dpi=dpi)
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 100)
    ax.set_aspect("equal")
    ax.axis("off")
    fig.patch.set_facecolor("white")
    return fig, ax


def _box(ax, x, y, w, h, label, *, fill=COLOR_BOX, edge=COLOR_BOX_EDGE,
         text_color="#111827", fontsize=10, fontweight="normal", rounded=True):
    """Draw a rounded rectangle with centred label."""
    style = "round,pad=0.02,rounding_size=0.5" if rounded else "square"
    box = FancyBboxPatch(
        (x, y), w, h,
        boxstyle=style, linewidth=1.2,
        edgecolor=edge, facecolor=fill,
    )
    ax.add_patch(box)
    ax.text(
        x + w / 2, y + h / 2, label,
        ha="center", va="center",
        fontsize=fontsize, color=text_color, fontweight=fontweight,
        wrap=True,
    )


def _arrow(ax, x1, y1, x2, y2, *, color="#374151", style="-|>",
           lw=1.2, label=None):
    """Draw an arrow with optional label near the midpoint."""
    arrow = FancyArrowPatch(
        (x1, y1), (x2, y2),
        arrowstyle=style, mutation_scale=12, linewidth=lw, color=color,
        connectionstyle="arc3,rad=0",
    )
    ax.add_patch(arrow)
    if label:
        ax.text(
            (x1 + x2) / 2, (y1 + y2) / 2 + 1.5, label,
            ha="center", va="bottom", fontsize=8, color=color,
            style="italic",
        )


def _save(fig, name: str):
    path = DIAGRAMS / f"{name}.png"
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print(f"  diagram → {path.relative_to(ROOT)}")
    return path


# ─── Diagrams ────────────────────────────────────────────────────────

def diagram_01_system_architecture():
    """High-level component diagram."""
    fig, ax = _new_canvas(14, 9)
    ax.text(50, 96, "System Architecture", ha="center", fontsize=14,
            fontweight="bold", color=COLOR_PRIMARY)

    # Client
    _box(ax, 35, 84, 30, 8, "Browser / API Client\n(MISA executive)",
         fill="#eff6ff", edge=COLOR_ACCENT, fontweight="bold")
    _arrow(ax, 50, 84, 50, 78, label="HTTPS")

    # Front Door (planned)
    _box(ax, 25, 68, 50, 8,
         "Azure Front Door + WAF  (planned)\nTLS · DDoS · Geo-fence · Rate limit",
         fill="#fef3c7", edge=COLOR_AMBER, fontsize=9)
    _arrow(ax, 50, 68, 50, 62)

    # FastAPI
    _box(ax, 15, 38, 70, 22, "", fill="#f3f4f6", edge="#9ca3af")
    ax.text(50, 58, "FastAPI Application (Uvicorn)", ha="center",
            fontsize=11, fontweight="bold", color=COLOR_PRIMARY)
    # Middleware
    _box(ax, 18, 50, 30, 5, "Middleware: CORS · Audit",
         fill="#ffffff", edge="#d1d5db", fontsize=8)
    _box(ax, 52, 50, 30, 5, "(planned) Rate limit · Injection",
         fill="#fef3c7", edge=COLOR_AMBER, fontsize=8)
    # Routers
    _box(ax, 18, 42, 20, 6, "/api/v1/chat", fill=COLOR_BOX, fontsize=9)
    _box(ax, 40, 42, 20, 6, "/api/v1/search", fill=COLOR_BOX, fontsize=9)
    _box(ax, 62, 42, 20, 6, "/api/v1/engagement",
         fill=COLOR_BOX, fontsize=9)
    _arrow(ax, 50, 42, 50, 32)

    # Service layer
    _box(ax, 10, 22, 80, 8,
         "Service Layer (24 modules)\n"
         "intent_router · depth_detector · correlator · curation · "
         "deep_profile · web_search · response_validator",
         fill="#e0e7ff", edge="#6366f1", fontsize=9)
    _arrow(ax, 30, 22, 18, 12)
    _arrow(ax, 70, 22, 82, 12)

    # Databases
    _box(ax, 5, 4, 25, 8, "PostgreSQL\n(local network)\n94 tables · read-only",
         fill="#dcfce7", edge=COLOR_GREEN, fontsize=9)
    _box(ax, 70, 4, 25, 8,
         "Azure OpenAI\n(Sweden Central)\ngpt-4.1-mini · store=false",
         fill="#dbeafe", edge=COLOR_ACCENT, fontsize=9)

    return _save(fig, "01_system_architecture")


def diagram_02_request_flow():
    """End-to-end chat turn lifecycle."""
    fig, ax = _new_canvas(14, 11)
    ax.text(50, 98, "End-to-End Chat Turn", ha="center", fontsize=14,
            fontweight="bold", color=COLOR_PRIMARY)

    steps = [
        ("1. POST /api/v1/chat",       "Browser · HTTP Basic auth",                "#eff6ff"),
        ("2. Middleware",              "CORS · Audit (mint correlation ID)",       "#fef3c7"),
        ("3. Intent classification",   "LLM call → 1 of 13 intents (~600ms)",      "#e0e7ff"),
        ("4. Depth detection",         "Regex → simple_fact / briefing / ... (<1ms)", "#e0e7ff"),
        ("5. Entity resolution",       "LLM + alias resolver + word-anchor ILIKE",  "#e0e7ff"),
        ("6. Direct-path correlator",  "Parallel SQL fan-out — 13 queries · ~15ms", "#dcfce7"),
        ("7. Privacy filter",          "_safe_row strips sensitive columns",        "#fee2e2"),
        ("8. Curation",                "LLM call with filtered payload (5-10s)",    "#dbeafe"),
        ("9. Output validator",        "Detect forbidden tokens · optional regen",  "#fef3c7"),
        ("10. Audit log emit",         "JSON line → audit.jsonl + stdout · SIEM",   "#fef3c7"),
        ("11. JSON response",          "{answer, trace, debug} → browser",          "#eff6ff"),
    ]
    y = 88
    for i, (title, detail, color) in enumerate(steps):
        _box(ax, 8, y, 84, 6.5, "",
             fill=color, edge="#9ca3af", rounded=True)
        ax.text(11, y + 4.5, title, ha="left", va="center",
                fontsize=10, fontweight="bold", color=COLOR_PRIMARY)
        ax.text(11, y + 1.8, detail, ha="left", va="center",
                fontsize=9, color="#374151")
        if i < len(steps) - 1:
            _arrow(ax, 50, y, 50, y - 1.2, lw=1.0)
        y -= 8

    return _save(fig, "02_request_flow")


def diagram_03_privacy_layers():
    """Defence-in-depth privacy filter layers."""
    fig, ax = _new_canvas(13, 9)
    ax.text(50, 96, "Defence-in-Depth Privacy Filtering", ha="center",
            fontsize=14, fontweight="bold", color=COLOR_PRIMARY)
    ax.text(50, 92,
            "Each layer is independent — failure of one does not expose data",
            ha="center", fontsize=10, color=COLOR_MUTED, style="italic")

    layers = [
        ("Postgres",
         "Layer 1: DB column deny list\n"
         "password / token / secret / api_key / ssn / private_key …",
         "#fef3c7", "#d97706"),
        ("In-app",
         "Layer 2: Sensitive-key substring filter (curation.py)\n"
         "review_status / reviewer_comments / linkedin / sec_filings …",
         "#e0e7ff", "#6366f1"),
        ("In-app",
         "Layer 3: Field-level PDPL redactor (database.py)\n"
         "Role-based: e.g. analyst role → personal_email = [REDACTED]",
         "#e0e7ff", "#6366f1"),
        ("In-app",
         "Layer 4: Truncation + null-stripping (_safe_row)\n"
         "1200-char cap on strings · drops empty values · recurses",
         "#e0e7ff", "#6366f1"),
        ("Post-LLM",
         "Layer 5: Output validator + scrubber (response_validator.py)\n"
         "Strips [web:N], (High), Source: DB, etc. before user sees it",
         "#dcfce7", "#2d6a4f"),
    ]
    y = 82
    for stage, text, fill, edge in layers:
        _box(ax, 5, y - 6, 12, 8, stage, fill="#ffffff",
             edge=COLOR_PRIMARY, fontsize=9, fontweight="bold")
        _box(ax, 20, y - 6, 75, 8, text, fill=fill, edge=edge, fontsize=9)
        y -= 13

    # Arrows
    for y_top in (75, 62, 49, 36):
        _arrow(ax, 57, y_top, 57, y_top - 3, lw=1.2)

    # End state
    _box(ax, 25, 7, 50, 6,
         "→  Azure OpenAI receives only privacy-filtered data  ←",
         fill="#dbeafe", edge=COLOR_ACCENT, fontweight="bold", fontsize=10)

    return _save(fig, "03_privacy_layers")


def diagram_04_network_topology():
    """Target production network architecture."""
    fig, ax = _new_canvas(14, 10)
    ax.text(50, 97, "Target Production Network Topology", ha="center",
            fontsize=14, fontweight="bold", color=COLOR_PRIMARY)

    # Client zone
    _box(ax, 5, 85, 20, 8,
         "MISA Executive\nCorporate net · MFA",
         fill="#eff6ff", edge=COLOR_ACCENT, fontsize=9, fontweight="bold")
    _arrow(ax, 15, 85, 15, 77, label="HTTPS")

    # Edge zone
    _box(ax, 5, 67, 90, 8,
         "Azure Front Door + WAF (Premium)\n"
         "TLS 1.2+ · OWASP rules · DDoS · Geo-fence (KSA) · Rate limit",
         fill="#fef3c7", edge=COLOR_AMBER, fontsize=10, fontweight="bold")
    _arrow(ax, 50, 67, 50, 60, label="Private Link")

    # App zone
    _box(ax, 10, 35, 80, 22, "", fill="#f9fafb", edge="#9ca3af")
    ax.text(50, 55, "Azure VNet — App Service Subnet", ha="center",
            fontsize=11, fontweight="bold", color=COLOR_PRIMARY)
    _box(ax, 15, 42, 30, 10,
         "App Service (Linux)\nFastAPI · Uvicorn\nManaged Identity",
         fill="#dbeafe", edge=COLOR_ACCENT, fontsize=9)
    _box(ax, 55, 42, 30, 10,
         "Audit log streaming\n→ Log Analytics\n→ Sentinel SIEM",
         fill="#fef3c7", edge=COLOR_AMBER, fontsize=9)
    _arrow(ax, 45, 47, 55, 47)

    # Data zone (private endpoints)
    _box(ax, 5, 12, 28, 18,
         "Postgres Flexible Server\n\nPrivate endpoint · TLS only\n"
         "PITR backups · Geo-redundant",
         fill="#dcfce7", edge=COLOR_GREEN, fontsize=9)
    _box(ax, 36, 12, 28, 18,
         "Azure Key Vault\n(HSM-backed)\n\nApp secrets · Cert mgmt\n"
         "Managed-identity access",
         fill="#fef3c7", edge=COLOR_AMBER, fontsize=9)
    _box(ax, 67, 12, 28, 18,
         "Azure OpenAI\n(Sweden Central / UAE N)\n\n"
         "Private endpoint · Enterprise DPA\nstore=false",
         fill="#dbeafe", edge=COLOR_ACCENT, fontsize=9)

    for x in (19, 50, 81):
        _arrow(ax, 30, 42, x, 30, color="#6b7280", lw=0.8)

    return _save(fig, "04_network_topology")


def diagram_05_data_flow_to_llm():
    """What crosses each boundary."""
    fig, ax = _new_canvas(14, 8)
    ax.text(50, 96, "What Crosses Each Network Boundary", ha="center",
            fontsize=14, fontweight="bold", color=COLOR_PRIMARY)

    # FastAPI center
    _box(ax, 40, 38, 20, 14,
         "FastAPI\nApplication",
         fill="#f3f4f6", edge=COLOR_PRIMARY,
         fontweight="bold", fontsize=11)

    # Postgres (left)
    _box(ax, 5, 38, 20, 14, "Postgres\n(local net)",
         fill="#dcfce7", edge=COLOR_GREEN, fontsize=10, fontweight="bold")
    _arrow(ax, 25, 45, 40, 45, label="SQL queries\n(parameterised)",
           color="#374151")
    _arrow(ax, 40, 41, 25, 41, label="Privacy-filtered rows",
           color="#374151")

    # Azure OpenAI (right top)
    _box(ax, 75, 60, 20, 14,
         "Azure OpenAI\n(Sweden C.)",
         fill="#dbeafe", edge=COLOR_ACCENT, fontsize=10, fontweight="bold")
    _arrow(ax, 60, 50, 75, 65,
           label="Question +\nFILTERED payload",
           color=COLOR_ACCENT)
    _arrow(ax, 75, 62, 60, 47, label="Markdown answer",
           color=COLOR_ACCENT)

    # Public OpenAI (right bottom — minimal)
    _box(ax, 75, 18, 20, 14,
         "OpenAI public\n(web search only)",
         fill="#fef3c7", edge=COLOR_AMBER, fontsize=10, fontweight="bold")
    _arrow(ax, 60, 42, 75, 25,
           label="Entity name ONLY",
           color=COLOR_AMBER)

    # Browser (top)
    _box(ax, 40, 78, 20, 12, "Browser",
         fill="#eff6ff", edge=COLOR_ACCENT, fontsize=10, fontweight="bold")
    _arrow(ax, 50, 78, 50, 52, color=COLOR_ACCENT, label="HTTPS request")
    _arrow(ax, 53, 52, 53, 78, color=COLOR_ACCENT, label="HTTPS response")

    # Legend
    ax.text(50, 6,
            "Privacy posture: DB rows never reach OpenAI public API. "
            "Only the entity name + user-typed text leaves the Azure "
            "tenant boundary.",
            ha="center", fontsize=9, color=COLOR_MUTED, style="italic",
            wrap=True)

    return _save(fig, "05_data_flow_to_llm")


def diagram_06_agentic_stack():
    """The 4-layer agentic analytics stack mapping."""
    fig, ax = _new_canvas(13, 9)
    ax.text(50, 96, "Agentic Analytics Stack — MISA Implementation",
            ha="center", fontsize=14, fontweight="bold", color=COLOR_PRIMARY)
    ax.text(50, 92,
            "Each layer addresses one or more of: Ambiguity · Staleness · Retrieval failure",
            ha="center", fontsize=9, color=COLOR_MUTED, style="italic")

    layers = [
        (78, "04 · Validation", "#e7d6c0",
         "response_validator · scrubber · missing-data line · "
         "anti-hallucination clauses · confidence · golden tests · feedback loop",
         "Ambiguity · Staleness · Retrieval"),
        (60, "03 · Skills", "#d97706",
         "intent_router (13 intents) · depth_detector · correlator · "
         "deep_profile · web_search · alias_resolver · llm_cache",
         "Retrieval · Staleness"),
        (42, "02 · Sources of truth", "#6a8c52",
         "company_profiles canonical · alias_resolver · country_resolver · "
         "word-anchored ILIKE · name-match guard",
         "Ambiguity"),
        (24, "01 · Data foundations", "#1f2937",
         "Postgres schema · 94 tables · FK relationships · "
         "(upstream MISA data warehouse owns freshness)",
         "Ambiguity · Staleness"),
    ]
    for y, title, color, modules, addresses in layers:
        text_color = "white" if color in ("#1f2937", "#d97706", "#6a8c52") else "#111827"
        _box(ax, 8, y, 84, 14, "",
             fill=color, edge="none", rounded=True)
        ax.text(12, y + 10, title, ha="left", va="center",
                fontsize=12, fontweight="bold", color=text_color)
        ax.text(12, y + 6.5, modules, ha="left", va="center",
                fontsize=8.5, color=text_color, wrap=True)
        ax.text(12, y + 3, f"Addresses: {addresses}",
                ha="left", va="center", fontsize=8, color=text_color,
                style="italic")

    return _save(fig, "06_agentic_stack")


# ─── Docx helpers ────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _add_page_number(footer_para):
    """Insert PAGE field into a footer paragraph."""
    run = footer_para.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _add_toc(doc):
    """Insert a Word native TOC field. Reader hits F9 (or Word auto-prompts)
    to populate. The field is the only way to get a true hyperlinked TOC
    that updates as headings change."""
    para = doc.add_paragraph()
    run = para.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = r'TOC \o "1-3" \h \z \u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and select 'Update Field' to populate this table of contents."
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(placeholder)
    run._r.append(fldChar3)


def _add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0a73f0")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_styled_table(doc, headers: list[str], rows: list[list[str]], *,
                       header_fill="#1a365d", header_color="FFFFFF",
                       col_widths_in: list[float] | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(header_color)
        run.font.size = Pt(10)
        _set_cell_bg(hdr[i], header_fill)

    # Body
    for ri, row in enumerate(rows, start=1):
        cells = table.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)

    # Col widths
    if col_widths_in:
        for col_idx, w in enumerate(col_widths_in):
            for row in table.rows:
                row.cells[col_idx].width = Inches(w)

    return table


def _heading(doc, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    return h


def _para(doc, text: str, *, bold=False, italic=False, size=11,
           color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    return p


def _bullets(doc, items: list[str]):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")


def _embed_image(doc, path: Path, *, width_in: float = 6.5,
                  caption: str | None = None):
    doc.add_picture(str(path), width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(9)
        cap_run.font.color.rgb = RGBColor.from_string("6B7280")


# ─── Build the document ──────────────────────────────────────────────

def build():
    print("\n▸ Generating diagrams...")
    img_arch = diagram_01_system_architecture()
    img_flow = diagram_02_request_flow()
    img_privacy = diagram_03_privacy_layers()
    img_network = diagram_04_network_topology()
    img_data = diagram_05_data_flow_to_llm()
    img_stack = diagram_06_agentic_stack()

    print("\n▸ Building document...")
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Heading colours
    for lvl in (1, 2, 3):
        h = doc.styles[f"Heading {lvl}"]
        h.font.color.rgb = RGBColor.from_string("1A365D")

    # Page setup + footer with page number
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("MISA Intelligence Chatbot — Design Document  ·  Page ")
    _add_page_number(fp)

    # ─── Cover page ──────────────────────────────────────────────
    for _ in range(3): doc.add_paragraph()
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover_title.add_run("MISA Intelligence Chatbot")
    r.font.size = Pt(28); r.bold = True
    r.font.color.rgb = RGBColor.from_string("1A365D")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("System Design Document")
    r.font.size = Pt(20); r.italic = True
    r.font.color.rgb = RGBColor.from_string("6B7280")

    for _ in range(6): doc.add_paragraph()

    meta = [
        ("Document version", "1.0"),
        ("Last updated", "2026-06-09"),
        ("Owner", "MISA Intelligence Engineering"),
        ("Audience", "Engineering · Security · Compliance (NDMO / NCA) · DevOps · MISA IT"),
        ("Classification", "Internal — Confidential"),
    ]
    t = doc.add_table(rows=len(meta), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta):
        c1, c2 = t.rows[i].cells
        c1.text = ""; c2.text = ""
        r1 = c1.paragraphs[0].add_run(k)
        r1.bold = True; r1.font.color.rgb = RGBColor.from_string("1A365D")
        c2.paragraphs[0].add_run(v)
        c1.width = Inches(2.0); c2.width = Inches(4.5)

    doc.add_page_break()

    # ─── TOC ─────────────────────────────────────────────────────
    _heading(doc, "Table of Contents", 1)
    _para(doc,
          "This is a Microsoft Word native table of contents. To populate "
          "it, click the placeholder line below, right-click, and select "
          "“Update Field” (or press F9). Word will auto-build the linked "
          "TOC from the headings in this document.",
          italic=True, color="6B7280", size=10)
    _add_toc(doc)
    doc.add_page_break()

    # ─── 1. Executive Summary ────────────────────────────────────
    _heading(doc, "1. Executive Summary", 1)
    _heading(doc, "1.1 What the System Is", 2)
    _para(doc,
          "The MISA Intelligence Chatbot is a natural-language Q&A system "
          "that lets MISA executives (Ministers, Deputy Ministers, sector "
          "leads, CIOs) ask plain-English questions about Saudi investment "
          "opportunities, foreign companies, executive contacts, and "
          "engagement history — and receive structured, actionable "
          "briefings sourced from MISA's internal database.")

    _heading(doc, "1.2 What It Is NOT", 2)
    _bullets(doc, [
        "Not a general-purpose ChatGPT clone — answers are constrained to MISA-internal data",
        "Not a search engine — answers are synthesised briefings, not link lists",
        "Not a public service — closed system for authenticated MISA staff",
        "Not a transactional system — read-only; never writes to the database",
    ])

    _heading(doc, "1.3 Core Design Principles", 2)
    _add_styled_table(doc,
        ["Principle", "How it's enforced"],
        [
            ["Privacy first",
             "DB rows routed through privacy filter before any external LLM call; sensitive column allowlist; deny list at SQL layer"],
            ["Quality through specialisation",
             "13 distinct intent paths instead of one mega-prompt; depth detector matches answer breadth to question depth"],
            ["Honest provenance",
             "No fabricated facts; explicit ‘Internal records do not currently show: X’ when data missing; web-grounded facts tagged [web], DB facts [DB], inferences [inferred]"],
            ["Reversible operations",
             "Every architectural change gated by env var; default behaviour preserved when flag is off"],
            ["Defence in depth",
             "Multiple independent privacy layers so a single layer's failure doesn't expose data"],
        ],
        col_widths_in=[1.8, 4.7],
    )

    _heading(doc, "1.4 Current Maturity", 2)
    _bullets(doc, [
        "Codebase: ~13,700 lines of Python, 30+ service modules",
        "Test coverage: 63 golden regression tests; live integration suite",
        "Production readiness: currently single-tenant local deployment; "
        "not yet hardened for MISA executive production use (see Section 16)",
    ])

    doc.add_page_break()

    # ─── 2. System Architecture ──────────────────────────────────
    _heading(doc, "2. System Architecture", 1)
    _heading(doc, "2.1 Component Diagram", 2)
    _embed_image(doc, img_arch,
                 caption="Figure 1 — High-level component diagram. Yellow boxes denote planned production additions; green/blue boxes are live today.")

    _heading(doc, "2.2 Technology Stack", 2)
    _add_styled_table(doc,
        ["Layer", "Technology", "Version", "Rationale"],
        [
            ["Runtime", "Python", "3.9+", "Mature async + DB ecosystem"],
            ["HTTP framework", "FastAPI", "0.115+", "Native async, OpenAPI, DI"],
            ["ASGI server", "Uvicorn", "0.34+", "Production-grade ASGI 3.0"],
            ["DB driver", "psycopg2-binary", "2.9+", "Battle-tested Postgres"],
            ["LLM SDK", "openai (Python)", "2.40+", "Official; speaks Azure + public"],
            ["Database", "PostgreSQL", "14+", "Existing MISA standard"],
            ["LLM provider", "Azure OpenAI", "API 2024-08-01-preview", "Data residency in Microsoft tenant under enterprise DPA"],
            ["PDF rendering", "WeasyPrint", "62+", "For /api/v1/export/pdf"],
        ],
        col_widths_in=[1.2, 1.6, 1.2, 2.5],
    )

    doc.add_page_break()

    # ─── 3. End-to-End Request Flow ─────────────────────────────
    _heading(doc, "3. End-to-End Request Flow", 1)
    _para(doc,
          "A single chat turn (\"Tell me about Apple's Saudi presence\") "
          "executes the 11 steps shown below. Typical wall-clock latency: "
          "5–10s for simple facts, 10–15s for executive briefings.")
    _embed_image(doc, img_flow,
                 caption="Figure 2 — End-to-end lifecycle of a single chat turn.")

    doc.add_page_break()

    # ─── 4. Database Layer ──────────────────────────────────────
    _heading(doc, "4. Database Layer", 1)
    _heading(doc, "4.1 Schema Families", 2)
    _para(doc,
          "The system queries 94 tables but only ~30 are actively used in "
          "the executive briefing path. The schema is organised into five "
          "families:")
    _add_styled_table(doc,
        ["Family", "Anchor table", "Joined tables (selected)"],
        [
            ["Company", "company_profiles",
             "executives, news, ai_insights, business_units, competitors, financial_performances, geographic_revenues, global_presences"],
            ["MISA Engagement", "(no single anchor)",
             "misa_contact_details, opportunities, strategic_investors, meetings, engagements, meeting_notes"],
            ["Country", "country_profiles",
             "key_indicators, infrastructures, policy_incentives, vision_outlooks, risk_stabilities, strategic_opportunities, trade_partners, top_commodities, free_zones, insights"],
            ["Regulatory", "rhq_licenses",
             "rhq_topexecutives, executives (Saudi directory)"],
            ["Reference", "fdi_data, sectors, programs",
             "Lookup / taxonomy"],
        ],
        col_widths_in=[1.3, 1.5, 3.7],
    )

    _heading(doc, "4.2 Query Patterns", 2)
    _bullets(doc, [
        "Parameterised queries via psycopg2 %s substitution — no string concatenation of user values into SQL (SQL injection defence)",
        "Word-anchored ILIKE using PostgreSQL \\m...\\M regex boundaries for fuzzy entity match",
        "pg_trgm trigram similarity for close-but-not-exact matches",
        "Alias resolver runs BEFORE any search — brand to legal name",
        "Cross-table correlation via parallel ThreadPoolExecutor — 13 queries in ~15ms (vs ~300ms serial)",
    ])

    _heading(doc, "4.3 Connection Posture", 2)
    _bullets(doc, [
        "Application is read-only — no INSERT/UPDATE/DELETE in the codebase",
        "DB role used by the app should be granted SELECT only (defence in depth)",
        "Connection string loaded from env at process start — not from request input",
        "Production: TLS sslmode=verify-full with CA pinning; private endpoint",
    ])

    doc.add_page_break()

    # ─── 5. Application Layer ──────────────────────────────────
    _heading(doc, "5. Application Layer", 1)
    _heading(doc, "5.1 Routers", 2)
    _add_styled_table(doc,
        ["Endpoint", "Purpose", "LLM calls / request"],
        [
            ["POST /api/v1/chat", "Main NL Q&A", "2–4 (intent · entity · curation · optional validator)"],
            ["POST /api/v1/search", "Structured filter lookup", "0"],
            ["POST /api/v1/engagement/generate", "Long-form investor dossier", "1 streaming with native web search"],
            ["POST /api/v1/feedback", "Thumbs up/down logging", "0"],
            ["POST /api/v1/export/pdf", "Convert answer to PDF", "0"],
            ["GET /health", "Liveness check", "0"],
        ],
        col_widths_in=[2.0, 2.0, 2.5],
    )

    _heading(doc, "5.2 Service Modules (selected)", 2)
    _add_styled_table(doc,
        ["Module", "Responsibility", "LOC"],
        [
            ["chat_engine.py", "Orchestrator — routes intent → direct path", "3,400"],
            ["intent_router.py", "LLM-based intent classification (13 intents)", "790"],
            ["correlator.py", "Parallel FK fan-out for company / person / country", "533"],
            ["curation.py", "Builds prompt; calls LLM; runs validator", "750"],
            ["deep_profile.py", "/profile mode — 3-pillar with web grounding", "720"],
            ["depth_detector.py", "Regex-based depth classification", "195"],
            ["response_validator.py", "Detects forbidden output strings; triggers regen", "280"],
            ["audit_log.py", "Audit middleware (NCA ECC 2-12)", "220"],
            ["alias_resolver.py", "Brand → legal-name mapping", "320"],
            ["country_resolver.py", "Country-name normalisation", "240"],
            ["web_search.py", "Wraps OpenAI search-preview model", "145"],
            ["llm_cache.py", "LRU+TTL cache for repeat LLM calls", "95"],
        ],
        col_widths_in=[1.7, 4.0, 0.8],
    )

    doc.add_page_break()

    # ─── 6. LLM Integration ─────────────────────────────────────
    _heading(doc, "6. LLM Integration", 1)
    _heading(doc, "6.1 Azure OpenAI (Primary)", 2)
    _add_styled_table(doc,
        ["Property", "Value"],
        [
            ["Resource", "<azure-openai-resource> (Azure OpenAI service)"],
            ["Region", "Sweden Central"],
            ["Endpoint", "https://<azure-openai-resource>.openai.azure.com"],
            ["Deployment name", "gpt-4.1-mini"],
            ["Model family", "gpt-4.1-mini (April 2025 release)"],
            ["API version", "2024-08-01-preview"],
            ["Retention", "store=false (no training, no retention)"],
            ["Data Processing Agreement", "Microsoft Enterprise DPA"],
        ],
        col_widths_in=[2.0, 4.5],
    )

    _heading(doc, "6.2 When Azure Is Used", 2)
    _add_styled_table(doc,
        ["Call site", "Purpose", "Sees DB data?"],
        [
            ["intent_router.classify_intent()", "Map question → intent label", "No"],
            ["chat_engine._extract_entity_from_question()", "Extract entity name", "No"],
            ["curation.curate_company_insights()", "Compose final answer", "YES (privacy-filtered)"],
            ["response_validator.validate_curation()", "Optional regen on style violations", "Yes"],
            ["deep_profile._compose_three_pillar()", "/profile mode composition", "Yes"],
        ],
        col_widths_in=[2.5, 2.5, 1.5],
    )

    _heading(doc, "6.3 When Public OpenAI Is Still Used", 2)
    _bullets(doc, [
        "Web search in /profile mode — uses gpt-4o-mini-search-preview, not yet on Azure. Only entity name leaves the boundary; no DB rows.",
        "/api/v1/engagements — uses gpt-4o with native web_search_preview tool. Only entity name + user-typed context; no DB rows.",
    ])

    doc.add_page_break()

    # ─── 7. Data Privacy & Information Flow ─────────────────────
    _heading(doc, "7. Data Privacy & Information Flow", 1)
    _heading(doc, "7.1 What Crosses Each Network Boundary", 2)
    _embed_image(doc, img_data,
                 caption="Figure 3 — Information flow across system boundaries.")

    _heading(doc, "7.2 Defence-in-Depth Privacy Filtering", 2)
    _embed_image(doc, img_privacy,
                 caption="Figure 4 — Five independent privacy layers. Each layer is sufficient on its own to prevent the specific class of leak it targets.")

    _heading(doc, "7.3 What is NEVER Sent to Any LLM", 2)
    _bullets(doc, [
        "DB passwords, API keys, tokens (Layer 1)",
        "Audit reviewer comments, internal team notes (Layer 2)",
        "External source paths (Factiva, LinkedIn URLs) (Layer 2)",
        "Environment variables, server paths, stack traces",
        "Other users' questions or answers (no cross-tenant data)",
    ])

    _heading(doc, "7.4 What IS Sent (Be Transparent)", 2)
    _para(doc,
          "For a typical chat turn about a company, Azure OpenAI receives:")
    _bullets(doc, [
        "The user's question (verbatim)",
        "The system prompt (static; describes style + format expectations)",
        "A JSON payload of privacy-filtered DB rows: company profile, executives, opportunities, MISA contacts (incl. names, emails, phones for non-redacted contacts), meeting agendas, meeting notes (incl. free-text), engagement action points",
    ])

    _para(doc, "Under Microsoft's Enterprise DPA:", bold=True)
    _bullets(doc, [
        "No training on this data",
        "No retention beyond the in-flight request (store=false)",
        "Region locked to Sweden Central",
        "Encryption in transit (TLS 1.2+) and at rest (AES-256)",
    ])

    doc.add_page_break()

    # ─── 8. Network Architecture ──────────────────────────────
    _heading(doc, "8. Network Architecture", 1)
    _heading(doc, "8.1 Target Production Topology", 2)
    _embed_image(doc, img_network,
                 caption="Figure 5 — Recommended production network topology. Front Door + WAF + private endpoints across the entire data path.")

    _heading(doc, "8.2 TLS / Cryptography Requirements", 2)
    _add_styled_table(doc,
        ["Layer", "Requirement"],
        [
            ["Transport", "TLS 1.2 minimum, TLS 1.3 preferred — enforced at Front Door"],
            ["Cipher suites", "Mandate forward secrecy (ECDHE); ban legacy (RC4, 3DES, MD5)"],
            ["HSTS", "1-year max-age + includeSubDomains"],
            ["At rest", "AES-256 (Azure-managed keys by default; CMK via Key Vault for higher tiers)"],
            ["Postgres link", "sslmode=verify-full with CA pinning"],
        ],
        col_widths_in=[1.6, 4.9],
    )

    doc.add_page_break()

    # ─── 9. Agentic Stack ─────────────────────────────────────
    _heading(doc, "9. Architecture Pattern — Agentic Analytics Stack", 1)
    _para(doc,
          "Our system is built in the canonical 4-layer agentic analytics "
          "pattern. Each layer addresses one or more of the three core "
          "failure modes — Ambiguity, Model Staleness, Retrieval Failure.")
    _embed_image(doc, img_stack,
                 caption="Figure 6 — Agentic stack with MISA-specific module mapping per layer.")

    doc.add_page_break()

    # ─── 10. Compliance Mapping ───────────────────────────────
    _heading(doc, "10. Compliance Mapping — NDMO + NCA", 1)
    _heading(doc, "10.1 NCA Essential Cybersecurity Controls (ECC)", 2)
    _add_styled_table(doc,
        ["Control", "Status"],
        [
            ["1-1 Cybersecurity Strategy", "🟠 Org work; this doc is an input"],
            ["2-2 Identity & Access Mgmt (MFA, RBAC)", "🔴 Gap — basic auth today; Entra ID + MFA planned"],
            ["2-7 Cryptography (TLS, at-rest, key mgmt)", "🟡 Partial; Azure inherits at rest"],
            ["2-8 Backup & Recovery", "🟠 Org work; Postgres backups inherited"],
            ["2-9 Vulnerability Management", "🔴 Gap (add scanning in CI)"],
            ["2-10 Penetration Testing", "🔴 Gap (engage NCA-licensed firm)"],
            ["2-11 Cybersecurity Event Logs", "✅ Implemented (audit_log middleware)"],
            ["2-12 Incident Management", "🔴 Gap (IR plan needed)"],
            ["2-13 Physical Security", "✅ Inherited from Azure"],
            ["2-14 Web Application Security (OWASP)", "🟡 Partial — needs WAF"],
            ["4-1 Third Party Cybersecurity (Vendor risk)", "🟠 Azure DPA pending MISA legal review"],
            ["5-1 CCC (Cloud Computing Controls)", "🟠 Sweden Central — verify SDAIA acceptance"],
        ],
        col_widths_in=[3.0, 3.5],
    )

    _heading(doc, "10.2 NDMO Data Management Framework", 2)
    _add_styled_table(doc,
        ["Domain", "Status"],
        [
            ["DG-3 Data Catalog & Metadata", "🟠 Schema exists; classification policy needed"],
            ["DG-7 Data Architecture", "✅ Documented (this doc)"],
            ["DG-10 Data Sharing & Interoperability (cross-border)", "🟠 Cross-border to Sweden Central"],
            ["DG-12 Data Security", "🟡 Partial — covered in §7 + §10"],
            ["DG-13 Personal Data Protection (PDPL)", "🔴 Gap — no DSR endpoints"],
        ],
        col_widths_in=[3.5, 3.0],
    )

    _heading(doc, "10.3 PDPL Article Mapping", 2)
    _add_styled_table(doc,
        ["Article", "Requirement", "Status"],
        [
            ["Art. 5", "Lawful basis for processing", "🟠 Documented basis needed"],
            ["Art. 10", "Data minimisation", "🟡 Privacy filter helps; can go further"],
            ["Art. 16-19", "Data Subject Rights", "🔴 No DSR endpoints"],
            ["Art. 25", "Records of Processing (ROPA)", "🟠 This doc is an input"],
            ["Art. 26", "DPO appointment", "🟠 Org decision"],
            ["Art. 29", "Cross-border transfer", "🟠 Sweden Central needs SDAIA approval"],
            ["Art. 33", "Incident notification (72h)", "🟠 IR plan needed"],
        ],
        col_widths_in=[1.0, 3.0, 2.5],
    )

    doc.add_page_break()

    # ─── 11. Production Readiness Roadmap ───────────────────────
    _heading(doc, "11. Production Readiness Roadmap", 1)
    _heading(doc, "11.1 Critical-Path (Must-Have Before Production)", 2)
    _add_styled_table(doc,
        ["#", "Item", "Effort", "Owner"],
        [
            ["1", "Replace Basic Auth with Entra ID + MFA", "2–3 days", "Backend + Azure"],
            ["2", "Implement RBAC (per-role data redaction)", "2 days", "Backend"],
            ["3", "Move secrets to Key Vault", "1 day", "Backend + Azure"],
            ["4", "Deploy Azure Front Door + WAF", "1 day", "DevOps"],
            ["5", "Rate limiting middleware", "0.5 day", "Backend"],
            ["6", "Prompt injection sanitiser", "1 day", "Backend"],
            ["7", "Penetration test by NCA-licensed firm", "1–2 weeks", "Procurement → external"],
            ["8", "Vulnerability scanning in CI", "0.5 day", "DevOps"],
            ["9", "TLS 1.2+ enforcement + HSTS", "0.5 day", "DevOps"],
            ["10", "DSR endpoints (/api/v1/dsr/*)", "1 day", "Backend"],
        ],
        col_widths_in=[0.4, 3.4, 1.0, 1.7],
    )
    _para(doc, "Estimated total: ~3–4 weeks engineering + parallel pentest.",
          italic=True, color="6B7280")

    _heading(doc, "11.2 Should-Have (First Quarter)", 2)
    _bullets(doc, [
        "Strict PII redactor (opt-in flag) — 1 day",
        "Output classification banner — 0.5 day",
        "Migrate Sweden Central → UAE North — 1 day after Azure provisioning",
        "SIEM integration with Azure Sentinel — 1–2 days",
        "Backup verification + DR runbook — 2–3 days",
        "Tabletop IR drill — 1 day",
    ])

    _heading(doc, "11.3 Could-Have (Strategic)", 2)
    _bullets(doc, [
        "Self-hosted LLM (Llama 3 / Mistral) for full data residency — 1–2 weeks",
        "Move to KSA Azure region (when GA) — 1–2 days",
        "Active-active multi-region deployment — 1–2 weeks",
        "Customer-managed encryption keys (BYOK) — 1 week",
    ])

    doc.add_page_break()

    # ─── 12. Known Risks & Limitations ─────────────────────────
    _heading(doc, "12. Known Risks & Limitations", 1)
    _heading(doc, "12.1 Quality Risks", 2)
    _add_styled_table(doc,
        ["Risk", "Mitigation"],
        [
            ["LLM hallucination of facts not in DB",
             "Anti-hallucination clauses in every prompt; output validator scrubs unsourced bullets; missing-data transparency directive"],
            ["Wrong entity resolved (e.g. ‘Apple’ → Mandiant Inc.)",
             "Alias resolver + word-anchored ILIKE + name-match filter"],
            ["Stale data",
             "DB is system of record; freshness depends on upstream pipelines"],
            ["Inconsistent prose style",
             "Single style_guide.py source of truth; output validator enforces"],
        ],
        col_widths_in=[2.5, 4.0],
    )

    _heading(doc, "12.2 Security Risks", 2)
    _add_styled_table(doc,
        ["Risk", "Severity", "Status"],
        [
            ["Basic auth with shared password", "🔴 Critical", "Production: Entra ID + MFA"],
            ["No RBAC", "🔴 Critical", "Production: role-based redaction"],
            ["Secrets in .env file", "🟠 High", "Production: Key Vault + managed identity"],
            ["No WAF", "🟠 High", "Production: Front Door Premium"],
            ["No rate limiting", "🟠 High", "Application middleware (planned)"],
            ["No prompt-injection defence", "🟠 High", "Sanitiser (planned)"],
            ["Cross-border transfer (Sweden Central)", "🟡 Medium", "Document for NDMO; migrate to UAE North"],
            ["No DSR endpoints", "🟡 Medium", "Implement /api/v1/dsr/*"],
            ["No penetration test", "🔴 Critical pre-prod", "Engage NCA-licensed firm"],
        ],
        col_widths_in=[3.0, 1.3, 2.2],
    )

    doc.add_page_break()

    # ─── Appendix A — Config Reference ────────────────────────
    _heading(doc, "Appendix A — Configuration Reference", 1)
    _para(doc,
          "All configuration via environment variables. See .env.example "
          "in the repository for the complete list. Key categories:")
    _add_styled_table(doc,
        ["Category", "Variables"],
        [
            ["Postgres", "PG_HOST · PG_PORT · PG_DB · PG_USER · PG_PASSWORD"],
            ["OpenAI (public)", "OPENAI_API_KEY · OPENAI_MODEL · OPENAI_MAX_RETRIES"],
            ["Azure OpenAI", "MISA_USE_AZURE_OPENAI · AZURE_OPENAI_ENDPOINT · AZURE_OPENAI_API_KEY · AZURE_OPENAI_API_VERSION · AZURE_OPENAI_DEPLOYMENT"],
            ["Engagement dossier", "MISA_ENGAGEMENT_OPENAI_API_KEY · MISA_ENGAGEMENT_OPENAI_MODEL"],
            ["Curation", "MISA_CHAT_CURATION · MISA_CHAT_FALLBACK · MISA_CHAT_OPENAI_STORE · MISA_CHAT_CURATION_MAX_ROWS"],
            ["Auth", "API_USERNAME · API_PASSWORD"],
            ["Audit", "MISA_AUDIT_LOG · MISA_AUDIT_LOG_FILE · MISA_AUDIT_LOG_STDOUT · MISA_AUDIT_LOG_MAX_MB · MISA_AUDIT_LOG_BACKUPS · MISA_TRUST_XFF"],
        ],
        col_widths_in=[1.5, 5.0],
    )

    doc.add_page_break()

    # ─── Appendix B — Glossary ─────────────────────────────────
    _heading(doc, "Appendix B — Glossary", 1)
    glossary = [
        ("AOAI", "Azure OpenAI"),
        ("CCC", "Cloud Computing Controls (NCA, 1:2020)"),
        ("DPA", "Data Processing Agreement"),
        ("DPO", "Data Protection Officer"),
        ("DSR", "Data Subject Rights (PDPL)"),
        ("ECC", "Essential Cybersecurity Controls (NCA)"),
        ("Entra ID", "Microsoft's identity provider (formerly Azure AD)"),
        ("FK", "Foreign Key"),
        ("HSM", "Hardware Security Module"),
        ("IdP", "Identity Provider"),
        ("JWT", "JSON Web Token"),
        ("MFA", "Multi-Factor Authentication"),
        ("NCA", "National Cybersecurity Authority (Saudi)"),
        ("NDMO", "National Data Management Office (Saudi)"),
        ("PDPL", "Personal Data Protection Law (Saudi)"),
        ("PITR", "Point-In-Time Recovery"),
        ("RBAC", "Role-Based Access Control"),
        ("RHQ", "Regional Headquarters"),
        ("ROPA", "Records of Processing Activities"),
        ("SDAIA", "Saudi Data and AI Authority"),
        ("SIEM", "Security Information and Event Management"),
        ("TPM", "Tokens Per Minute"),
        ("WAF", "Web Application Firewall"),
    ]
    _add_styled_table(doc,
        ["Term", "Meaning"],
        [[t, m] for t, m in glossary],
        col_widths_in=[1.3, 5.2],
    )

    doc.add_page_break()
    _para(doc,
          "End of document. Maintained at docs/DESIGN.md (source) and "
          "regenerated to docs/MISA_Chatbot_Design_Document.docx via "
          "scripts/build_design_docx.py.",
          italic=True, color="6B7280", size=10)

    out = DOCS / "MISA_Chatbot_Design_Document.docx"
    doc.save(out)
    print(f"\n✓ Saved: {out.relative_to(ROOT)}")
    print(f"  size: {out.stat().st_size / 1024:.1f} KB")
    return out


if __name__ == "__main__":
    build()
